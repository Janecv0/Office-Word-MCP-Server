"""
Document creation and manipulation tools for Word Document Server.
"""
import os
import json
from typing import Dict, List, Optional, Any
from docx import Document

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension, create_document_copy
from word_document_server.utils.document_utils import get_document_properties, extract_document_text, get_document_structure, get_document_xml, insert_header_near_text, insert_line_or_paragraph_near_text
from word_document_server.core.styles import ensure_heading_style, ensure_table_style
from word_document_server.security.http_auth import build_download_url


def _resolve_output_target(filename: str) -> str:
    """Resolve target path into DOC_OUTPUT_DIR for basename-only inputs."""
    filename = ensure_docx_extension(filename)
    output_dir = os.getenv("DOC_OUTPUT_DIR")
    if output_dir and not os.path.isabs(filename) and os.path.dirname(filename) == "":
        os.makedirs(output_dir, exist_ok=True)
        return os.path.join(output_dir, filename)
    return filename


def _resolve_source_document(filename: str) -> Optional[str]:
    """Resolve a source document path, falling back to DOC_OUTPUT_DIR for basenames."""
    filename = ensure_docx_extension(filename)
    if os.path.exists(filename):
        return filename

    output_dir = os.getenv("DOC_OUTPUT_DIR")
    if output_dir and not os.path.isabs(filename):
        candidate = os.path.join(output_dir, os.path.basename(filename))
        if os.path.exists(candidate):
            return candidate

    return None


async def create_document(filename: str, title: Optional[str] = None, author: Optional[str] = None) -> str:
    """Create a new Word document with optional metadata.
    
    Args:
        filename: Name of the document to create (with or without .docx extension)
        title: Optional title for the document metadata
        author: Optional author for the document metadata
    """
    filename = _resolve_output_target(filename)
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot create document: {error_message}"
    
    try:
        doc = Document()
        
        # Set properties if provided
        if title:
            doc.core_properties.title = title
        if author:
            doc.core_properties.author = author
        
        # Ensure necessary styles exist
        ensure_heading_style(doc)
        ensure_table_style(doc)
        
        # Save the document
        doc.save(filename)
        
        return f"Document {filename} created successfully"
    except Exception as e:
        return f"Failed to create document: {str(e)}"


async def get_document_info(filename: str) -> str:
    """Get information about a Word document.
    
    Args:
        filename: Path to the Word document
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    try:
        properties = get_document_properties(filename)
        return json.dumps(properties, indent=2)
    except Exception as e:
        return f"Failed to get document info: {str(e)}"


async def get_document_text(filename: str) -> str:
    """Extract all text from a Word document.
    
    Args:
        filename: Path to the Word document
    """
    filename = ensure_docx_extension(filename)
    
    return extract_document_text(filename)


async def get_document_outline(filename: str) -> str:
    """Get the structure of a Word document.
    
    Args:
        filename: Path to the Word document
    """
    filename = ensure_docx_extension(filename)
    
    structure = get_document_structure(filename)
    return json.dumps(structure, indent=2)


async def list_available_documents(directory: str = ".") -> str:
    """List all .docx files in the specified directory.
    
    Args:
        directory: Directory to search for Word documents
    """
    try:
        output_dir = os.getenv("DOC_OUTPUT_DIR")
        if directory in (".", "") and output_dir:
            directory = output_dir

        if not os.path.exists(directory):
            return f"Directory {directory} does not exist"
        
        docx_files = [f for f in os.listdir(directory) if f.endswith('.docx')]
        
        if not docx_files:
            return f"No Word documents found in {directory}"
        
        result = f"Found {len(docx_files)} Word documents in {directory}:\n"
        for file in docx_files:
            file_path = os.path.join(directory, file)
            size = os.path.getsize(file_path) / 1024  # KB
            result += f"- {file} ({size:.2f} KB)\n"
        
        return result
    except Exception as e:
        return f"Failed to list documents: {str(e)}"


async def copy_document(source_filename: str, destination_filename: Optional[str] = None) -> str:
    """Create a copy of a Word document.
    
    Args:
        source_filename: Path to the source document
        destination_filename: Optional path for the copy. If not provided, a default name will be generated.
    """
    source_filename = ensure_docx_extension(source_filename)
    
    if destination_filename:
        destination_filename = ensure_docx_extension(destination_filename)
    
    success, message, new_path = create_document_copy(source_filename, destination_filename)
    if success:
        return message
    else:
        return f"Failed to copy document: {message}"


async def merge_documents(target_filename: str, source_filenames: List[str], add_page_breaks: bool = True) -> str:
    """Merge multiple Word documents into a single document.
    
    Args:
        target_filename: Path to the target document (will be created or overwritten)
        source_filenames: List of paths to source documents to merge
        add_page_breaks: If True, add page breaks between documents
    """
    from word_document_server.core.tables import copy_table
    
    target_filename = ensure_docx_extension(target_filename)
    
    # Check if target file is writeable
    is_writeable, error_message = check_file_writeable(target_filename)
    if not is_writeable:
        return f"Cannot create target document: {error_message}"
    
    # Validate all source documents exist
    missing_files = []
    for filename in source_filenames:
        doc_filename = ensure_docx_extension(filename)
        if not os.path.exists(doc_filename):
            missing_files.append(doc_filename)
    
    if missing_files:
        return f"Cannot merge documents. The following source files do not exist: {', '.join(missing_files)}"
    
    try:
        # Create a new document for the merged result
        target_doc = Document()
        
        # Process each source document
        for i, filename in enumerate(source_filenames):
            doc_filename = ensure_docx_extension(filename)
            source_doc = Document(doc_filename)
            
            # Add page break between documents (except before the first one)
            if add_page_breaks and i > 0:
                target_doc.add_page_break()
            
            # Copy all paragraphs
            for paragraph in source_doc.paragraphs:
                # Create a new paragraph with the same text and style
                new_paragraph = target_doc.add_paragraph(paragraph.text)
                new_paragraph.style = target_doc.styles['Normal']  # Default style
                
                # Try to match the style if possible
                try:
                    if paragraph.style and paragraph.style.name in target_doc.styles:
                        new_paragraph.style = target_doc.styles[paragraph.style.name]
                except:
                    pass
                
                # Copy run formatting
                for i, run in enumerate(paragraph.runs):
                    if i < len(new_paragraph.runs):
                        new_run = new_paragraph.runs[i]
                        # Copy basic formatting
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        # Font size if specified
                        if run.font.size:
                            new_run.font.size = run.font.size
            
            # Copy all tables
            for table in source_doc.tables:
                copy_table(table, target_doc)
        
        # Save the merged document
        target_doc.save(target_filename)
        return f"Successfully merged {len(source_filenames)} documents into {target_filename}"
    except Exception as e:
        return f"Failed to merge documents: {str(e)}"


async def get_document_xml_tool(filename: str) -> str:
    """Get the raw XML structure of a Word document."""
    return get_document_xml(filename)


async def save_document(file_path: str, source_filename: str) -> Dict[str, Any]:
    """Save a document copy to a target path, optionally constrained to DOC_OUTPUT_DIR.

    Args:
        file_path: Requested output path or filename
        source_filename: Existing source .docx to save/copy
    """
    source_filename = ensure_docx_extension(source_filename)
    source_path = _resolve_source_document(source_filename)
    if not source_path:
        checked = [source_filename]
        output_dir = os.getenv("DOC_OUTPUT_DIR")
        if output_dir and not os.path.isabs(source_filename):
            checked.append(os.path.join(output_dir, os.path.basename(source_filename)))
        return {"error": f"Document {source_filename} does not exist (checked: {', '.join(checked)})"}

    output_dir = os.getenv("DOC_OUTPUT_DIR")
    filename = os.path.basename(file_path) if file_path else ""

    try:
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
            if not filename:
                filename = os.path.basename(source_filename) or "document.docx"
            if not filename.lower().endswith(".docx"):
                filename += ".docx"
            save_path = os.path.join(output_dir, filename)
        else:
            if not file_path:
                return {"error": "file_path is required when DOC_OUTPUT_DIR is not configured"}
            save_path = ensure_docx_extension(file_path)
            filename = os.path.basename(save_path)

        if os.path.abspath(source_path) == os.path.abspath(save_path):
            result: Dict[str, Any] = {
                "message": f"Document already saved at {save_path}",
                "file_path": save_path,
            }
            download_base_url = os.getenv("DOC_DOWNLOAD_BASE_URL", os.getenv("MCP_DOWNLOAD_BASE_URL"))
            if download_base_url:
                result["download_url"] = build_download_url(download_base_url, filename)
            return result

        is_writeable, error_message = check_file_writeable(save_path)
        if not is_writeable:
            return {"error": f"Cannot save document: {error_message}"}

        success, message, saved_path = create_document_copy(source_path, save_path)
        if not success or not saved_path:
            return {"error": f"Failed to save document: {message}"}

        if not os.path.exists(saved_path):
            return {"error": f"Save reported success but file not found at {saved_path}"}

        result: Dict[str, Any] = {
            "message": f"Document saved to {saved_path}",
            "file_path": saved_path,
            "file_size_bytes": os.path.getsize(saved_path),
        }

        download_base_url = os.getenv("DOC_DOWNLOAD_BASE_URL", os.getenv("MCP_DOWNLOAD_BASE_URL"))
        if download_base_url:
            result["download_url"] = build_download_url(download_base_url, filename)

        return result
    except Exception as e:
        return {"error": f"Failed to save document: {str(e)}"}



